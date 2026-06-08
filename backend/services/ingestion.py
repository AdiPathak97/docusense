"""
Document ingestion pipeline: parse → chunk → embed → upsert.

Chunk size (500 tokens, overlap 50) is set in config.
If you change it, increment EMBEDDING_VERSION and re-embed all existing documents.
"""

import logging
import uuid
from pathlib import Path

from langchain.text_splitter import RecursiveCharacterTextSplitter

from backend.exceptions import IngestionError, LLMError
from backend.services.llm_provider import EmbeddingProvider
from backend.services.vector_store import VectorStoreClient

logger = logging.getLogger(__name__)

CHUNK_SIZE = 500    # tokens
CHUNK_OVERLAP = 50  # tokens
EMBEDDING_VERSION = 1


def parse_document(file_path: Path, content_type: str) -> list[dict]:
    """
    Parse uploaded file into a list of {page_number, text} dicts.
    Supported: application/pdf, application/vnd.openxmlformats-officedocument.wordprocessingml.document, text/plain
    """
    logger.debug(
        "parse_document — path=%s content_type=%s", file_path.name, content_type
    )
    try:
        if content_type == "application/pdf":
            import pypdfium2
            pages = []
            pdf = pypdfium2.PdfDocument(str(file_path))
            for i, page in enumerate(pdf, start=1):
                textpage = page.get_textpage()
                text = textpage.get_text_range() or ""
                if text.strip():
                    pages.append({"page_number": i, "text": text})
            logger.debug(
                "parse_document PDF ok — pages_with_text=%d", len(pages)
            )
            return pages

        if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            import docx
            doc = docx.Document(str(file_path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            logger.debug("parse_document DOCX ok — paragraphs extracted")
            return [{"page_number": 1, "text": text}]

        # text/plain
        text = file_path.read_text(encoding="utf-8", errors="replace")
        logger.debug("parse_document TXT ok — chars=%d", len(text))
        return [{"page_number": 1, "text": text}]

    except Exception as exc:
        raise IngestionError(
            f"Failed to parse document '{file_path.name}': {exc}",
        ) from exc


def chunk_pages(pages: list[dict]) -> list[dict]:
    """
    Split pages into chunks. Returns list of {chunk_index, page_number, text}.
    Uses RecursiveCharacterTextSplitter with tiktoken encoder.
    """
    logger.debug("chunk_pages — input_pages=%d", len(pages))
    try:
        splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
        )
        chunks = []
        chunk_index = 0
        for page in pages:
            splits = splitter.split_text(page["text"])
            for split in splits:
                chunks.append({
                    "chunk_index": chunk_index,
                    "page_number": page["page_number"],
                    "text": split,
                })
                chunk_index += 1
        logger.debug("chunk_pages ok — output_chunks=%d", len(chunks))
        return chunks
    except Exception as exc:
        raise IngestionError(f"Failed to chunk document: {exc}") from exc


async def ingest(
    document_id: str,
    document_name: str,
    file_path: Path,
    content_type: str,
    embedding_provider: EmbeddingProvider,
    vector_store: VectorStoreClient,
) -> list[dict]:
    """
    Full pipeline. Returns list of ingested chunks (id, chunk_index, page_number).
    Called from POST /api/documents/upload after DB record is created.
    """
    logger.info(
        "ingest start — document_id=%s name=%s", document_id, document_name
    )

    # parse and chunk both raise IngestionError on failure
    pages = parse_document(file_path, content_type)
    chunks = chunk_pages(pages)

    logger.debug(
        "ingest embedding — document_id=%s chunks=%d", document_id, len(chunks)
    )

    chunk_ids, embeddings, documents, metadatas = [], [], [], []
    for i, chunk in enumerate(chunks):
        chunk_id = str(uuid.uuid4())
        try:
            embedding = await embedding_provider.embed(chunk["text"])
        except LLMError:
            raise  # already typed; propagate as-is so caller sees LLMError
        except Exception as exc:
            raise IngestionError(
                f"Embedding failed for chunk {i} of document '{document_name}': {exc}",
                document_id=document_id,
            ) from exc

        chunk_ids.append(chunk_id)
        embeddings.append(embedding)
        documents.append(chunk["text"])
        metadatas.append({
            "document_id": document_id,
            "document_name": document_name,
            "page_number": chunk["page_number"],
            "chunk_index": chunk["chunk_index"],
            "embedding_version": EMBEDDING_VERSION,
        })

    # upsert_chunks raises VectorStoreError — let it propagate; the caller
    # (_run_ingestion in api/documents.py) catches all exceptions and marks
    # the document as failed.
    await vector_store.upsert_chunks(
        document_id, chunk_ids, embeddings, documents, metadatas
    )

    logger.info(
        "ingest complete — document_id=%s chunks=%d", document_id, len(chunks)
    )
    return [
        {"id": cid, "chunk_index": c["chunk_index"], "page_number": c["page_number"]}
        for cid, c in zip(chunk_ids, chunks)
    ]
